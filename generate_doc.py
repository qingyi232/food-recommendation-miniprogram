# -*- coding: utf-8 -*-
"""
生成《基于微信小程序的校园周边美食推荐系统》完整技术文档（Word格式）
三线表 + 论文风格
"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

doc = Document()

# ======================== 全局样式设置 ========================
style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(6)

# 标题样式
for level in range(1, 5):
    heading_style = doc.styles[f'Heading {level}']
    heading_font = heading_style.font
    heading_font.color.rgb = RGBColor(0, 0, 0)
    heading_style.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    if level == 1:
        heading_font.size = Pt(18)
        heading_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        heading_style.paragraph_format.space_before = Pt(24)
        heading_style.paragraph_format.space_after = Pt(18)
    elif level == 2:
        heading_font.size = Pt(16)
        heading_style.paragraph_format.space_before = Pt(18)
        heading_style.paragraph_format.space_after = Pt(12)
    elif level == 3:
        heading_font.size = Pt(14)
        heading_style.paragraph_format.space_before = Pt(12)
        heading_style.paragraph_format.space_after = Pt(8)
    else:
        heading_font.size = Pt(12)


def add_paragraph(text, bold=False, align=None, font_size=None, font_name=None, first_line_indent=True):
    """添加段落"""
    p = doc.add_paragraph()
    if first_line_indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    if align:
        p.paragraph_format.alignment = align
    run = p.add_run(text)
    if bold:
        run.bold = True
    if font_size:
        run.font.size = Pt(font_size)
    if font_name:
        run.font.name = font_name
        run.element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    return p


def make_three_line_table(headers, rows, col_widths=None):
    """创建三线表"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # 清除所有边框
    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = parse_xml(
                '<w:tcBorders %s>'
                '  <w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                '  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                '  <w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                '  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                '</w:tcBorders>' % nsdecls('w')
            )
            tcPr.append(tcBorders)

    # 表头行 - 顶线(粗) + 底线(细)
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        run.bold = True
        run.font.size = Pt(10.5)
        run.font.name = '黑体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = parse_xml(
            '<w:tcBorders %s>'
            '  <w:top w:val="single" w:sz="12" w:space="0" w:color="000000"/>'
            '  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            '  <w:bottom w:val="single" w:sz="6" w:space="0" w:color="000000"/>'
            '  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            '</w:tcBorders>' % nsdecls('w')
        )
        tcPr.append(tcBorders)

    # 数据行
    for row_idx, row_data in enumerate(rows):
        for col_idx, cell_text in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(cell_text))
            run.font.size = Pt(10.5)
            run.font.name = '宋体'
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

            # 最后一行加底线(粗)
            if row_idx == len(rows) - 1:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcBorders = parse_xml(
                    '<w:tcBorders %s>'
                    '  <w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                    '  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                    '  <w:bottom w:val="single" w:sz="12" w:space="0" w:color="000000"/>'
                    '  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                    '</w:tcBorders>' % nsdecls('w')
                )
                tcPr.append(tcBorders)

    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)

    return table


def add_table_caption(text):
    """表标题"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(10.5)
    run.font.name = '黑体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')


def add_code_block(code, caption=None):
    """添加代码块（灰底 Consolas 字体）"""
    if caption:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.first_line_indent = Cm(0.74)
        run = p.add_run(caption)
        run.bold = True
        run.font.size = Pt(10.5)
        run.font.name = '黑体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.rows[0].cells[0]
    cell.text = ''
    # 灰色背景
    shading = parse_xml('<w:shd {} w:fill="F5F5F5"/>'.format(nsdecls('w')))
    cell._tc.get_or_add_tcPr().append(shading)
    # 边框
    tcBorders = parse_xml(
        '<w:tcBorders %s>'
        '  <w:top w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        '  <w:left w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        '  <w:right w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        '</w:tcBorders>' % nsdecls('w')
    )
    cell._tc.get_or_add_tcPr().append(tcBorders)

    for i, line in enumerate(code.strip().split('\n')):
        if i == 0:
            p = cell.paragraphs[0]
        else:
            p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.first_line_indent = None
        run = p.add_run(line)
        run.font.size = Pt(9)
        run.font.name = 'Consolas'
        run.font.color.rgb = RGBColor(51, 51, 51)


# ======================== 封面 ========================
for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('基于微信小程序的校园周边美食推荐系统')
run.bold = True
run.font.size = Pt(26)
run.font.name = '黑体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('技  术  文  档')
run.bold = True
run.font.size = Pt(22)
run.font.name = '黑体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

for _ in range(6):
    doc.add_paragraph()

info_lines = [
    '项目名称：基于微信小程序的校园周边美食推荐系统',
    '技术栈：Spring Boot + MyBatis-Plus + 微信小程序 + MySQL',
    '文档版本：V1.0',
    '编写日期：2026年3月'
]
for line in info_lines:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(line)
    run.font.size = Pt(14)
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

doc.add_page_break()

# ======================== 目录页 ========================
doc.add_heading('目  录', level=1)
toc_items = [
    ('第一章  绪论', '1'),
    ('  1.1 项目背景', '1'),
    ('  1.2 研究目的与意义', '1'),
    ('  1.3 系统主要功能', '2'),
    ('第二章  相关技术介绍', '3'),
    ('  2.1 微信小程序开发技术', '3'),
    ('  2.2 Spring Boot框架', '3'),
    ('  2.3 MyBatis-Plus持久层框架', '4'),
    ('  2.4 MySQL数据库', '4'),
    ('  2.5 JWT身份认证', '4'),
    ('  2.6 协同过滤推荐算法', '5'),
    ('  2.7 其他技术组件', '5'),
    ('第三章  系统需求分析', '6'),
    ('  3.1 功能性需求分析', '6'),
    ('  3.2 非功能性需求分析', '7'),
    ('  3.3 系统用例分析', '7'),
    ('第四章  系统总体设计', '8'),
    ('  4.1 系统架构设计', '8'),
    ('  4.2 系统功能模块设计', '8'),
    ('  4.3 系统技术架构', '9'),
    ('第五章  数据库设计', '10'),
    ('  5.1 数据库概述', '10'),
    ('  5.2 数据表结构设计', '10'),
    ('  5.3 数据表关系说明', '18'),
    ('第六章  系统详细设计与实现', '19'),
    ('  6.1 后端接口设计', '19'),
    ('  6.2 用户模块实现', '20'),
    ('  6.3 商家模块实现', '21'),
    ('  6.4 店铺与菜品模块实现', '22'),
    ('  6.5 订单模块实现', '23'),
    ('  6.6 评价模块实现', '24'),
    ('  6.7 收藏与分享模块实现', '25'),
    ('  6.8 个性化推荐模块实现', '25'),
    ('  6.9 管理员后台模块实现', '27'),
    ('  6.10 系统安全与认证实现', '28'),
    ('第七章  系统前端设计与实现', '29'),
    ('  7.1 小程序页面结构', '29'),
    ('  7.2 网络请求封装', '30'),
    ('  7.3 管理后台前端', '30'),
    ('第八章  系统部署与运行', '31'),
    ('  8.1 开发环境', '31'),
    ('  8.2 部署步骤', '31'),
    ('第九章  总结与展望', '32'),
]
for item, page in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = None
    if item.startswith('  '):
        run = p.add_run('    ' + item.strip())
    else:
        run = p.add_run(item)
        run.bold = True
    run.font.size = Pt(12)
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

doc.add_page_break()

# ======================== 第一章 绪论 ========================
doc.add_heading('第一章  绪论', level=1)

doc.add_heading('1.1 项目背景', level=2)
add_paragraph(
    '随着移动互联网技术的快速发展和智能手机的广泛普及，微信小程序作为一种轻量级应用形态，'
    '凭借其"即用即走"的特性在各行各业得到了广泛应用。与此同时，高校校园周边的餐饮行业蓬勃发展，'
    '但学生在面对众多餐饮选择时往往缺乏有效的信息获取渠道，难以快速找到符合自身口味偏好的美食。'
    '传统的美食推荐方式主要依赖口碑传播和大众点评等平台，但这些方式缺乏针对校园用户群体的个性化推荐能力。'
)
add_paragraph(
    '在此背景下，本项目设计并实现了一套基于微信小程序的校园周边美食推荐系统。该系统综合运用协同过滤推荐算法'
    '和基于内容的推荐策略，通过采集和分析用户的浏览、点击、收藏、下单、评价等行为数据，'
    '为每位用户提供个性化的美食推荐服务。系统采用前后端分离架构，后端基于Spring Boot框架开发RESTful API，'
    '前端基于微信小程序原生开发框架构建，同时提供Web管理后台供管理员进行系统管理。'
)

doc.add_heading('1.2 研究目的与意义', level=2)
add_paragraph(
    '本项目的主要研究目的包括以下几个方面：'
)
add_paragraph('（1）设计并实现一个功能完善的校园周边美食推荐系统，满足学生查找、筛选、评价美食的需求。')
add_paragraph('（2）探索协同过滤推荐算法在美食推荐场景中的应用，通过用户行为分析提供精准的个性化推荐。')
add_paragraph('（3）实践前后端分离的微信小程序开发模式，积累Web应用全栈开发经验。')
add_paragraph('（4）构建用户、商家、管理员三端协同的完整业务生态系统。')
add_paragraph(
    '本项目的意义在于：一方面为校园用户提供便捷高效的美食发现途径，提升用餐体验；'
    '另一方面为校园周边商家提供数字化经营工具，帮助其拓展客源、提升服务质量；'
    '同时为管理员提供数据可视化管理平台，实现对整个平台的高效监管。'
)

doc.add_heading('1.3 系统主要功能', level=2)
add_paragraph('本系统主要包含以下核心功能模块：')

add_table_caption('表1-1 系统主要功能模块概览')
make_three_line_table(
    ['模块', '功能描述', '使用角色'],
    [
        ['用户管理', '注册、登录（含微信登录）、个人信息管理、口味偏好设置', '用户'],
        ['店铺浏览', '店铺列表、店铺详情、分类筛选、关键词搜索', '用户'],
        ['菜品浏览', '菜品列表、菜品详情、按店铺/分类查看', '用户'],
        ['个性化推荐', '基于协同过滤和内容的混合推荐算法，推荐店铺和菜品', '用户'],
        ['订单管理', '创建订单、查看订单列表、订单状态跟踪', '用户/商家'],
        ['评价系统', '发布评价（含图片）、多维度评分、商家回复', '用户/商家'],
        ['收藏功能', '收藏/取消收藏店铺和菜品', '用户'],
        ['分享功能', '分享店铺和菜品给好友', '用户'],
        ['商家管理', '商家注册与登录、店铺管理、菜品管理、订单处理、经营统计', '商家'],
        ['管理后台', '数据统计、用户管理、商家审核、评价管理、违规处理、数据可视化', '管理员'],
    ]
)

doc.add_page_break()

# ======================== 第二章 相关技术介绍 ========================
doc.add_heading('第二章  相关技术介绍', level=1)

doc.add_heading('2.1 微信小程序开发技术', level=2)
add_paragraph(
    '微信小程序是一种不需要下载安装即可使用的应用，用户通过微信扫一扫或搜索即可打开应用。'
    '小程序基于微信客户端运行，具有加载速度快、占用资源少、跨平台兼容等优点。'
    '本系统采用微信小程序原生开发框架进行前端开发，使用WXML（WeiXin Markup Language）进行页面结构描述，'
    'WXSS（WeiXin Style Sheets）进行样式控制，JavaScript进行逻辑处理。'
)
add_paragraph(
    '微信小程序的核心框架包括视图层（View）和逻辑层（App Service）两部分。'
    '视图层负责页面渲染，逻辑层负责处理业务逻辑和数据管理。两层之间通过数据绑定和事件系统进行通信。'
    '小程序还提供了丰富的原生API，包括网络请求（wx.request）、文件上传（wx.uploadFile）、'
    '本地存储（wx.setStorageSync）、页面导航（wx.navigateTo）等，为开发者提供了强大的能力支持。'
)

doc.add_heading('2.2 Spring Boot框架', level=2)
add_paragraph(
    'Spring Boot是由Pivotal团队开发的基于Spring框架的快速开发工具，旨在简化Spring应用的创建和部署过程。'
    'Spring Boot通过自动配置、起步依赖和内嵌服务器等特性，极大地降低了Spring应用的开发门槛。'
    '本系统后端采用Spring Boot 2.7.18版本进行开发。'
)
add_paragraph(
    'Spring Boot的主要优势包括：（1）自动配置（Auto-Configuration），根据项目依赖自动完成框架配置；'
    '（2）起步依赖（Starter Dependencies），通过引入特定的starter即可获得一组相关依赖；'
    '（3）内嵌Web服务器，支持Tomcat、Jetty等，无需外部部署；'
    '（4）Actuator监控端点，提供应用运行时的健康检查和指标监控能力。'
    '本系统主要使用了spring-boot-starter-web实现RESTful API开发，配合拦截器实现权限控制。'
)

doc.add_heading('2.3 MyBatis-Plus持久层框架', level=2)
add_paragraph(
    'MyBatis-Plus（简称MP）是MyBatis的增强工具，在MyBatis的基础上只做增强不做改变，旨在简化开发、提高效率。'
    '本系统采用MyBatis-Plus 3.5.3.1版本作为数据持久层框架。MP提供了强大的CRUD操作能力，'
    '开发者无需编写大量XML映射文件即可完成常见的数据库操作。'
)
add_paragraph(
    '本系统中使用了MyBatis-Plus的以下核心特性：'
    '（1）BaseMapper通用CRUD接口，提供selectById、insert、updateById、deleteById等方法；'
    '（2）LambdaQueryWrapper条件构造器，支持Lambda表达式构建类型安全的查询条件；'
    '（3）ServiceImpl通用Service实现类，提供save、page、list等高级方法；'
    '（4）分页插件（PaginationInnerInterceptor），实现物理分页查询；'
    '（5）自动填充功能（MetaObjectHandler），自动填充createTime和updateTime字段。'
)

doc.add_heading('2.4 MySQL数据库', level=2)
add_paragraph(
    'MySQL是一款开源的关系型数据库管理系统，以其高性能、高可靠性和易用性著称。'
    '本系统采用MySQL 8.0版本作为数据存储引擎，使用InnoDB存储引擎以支持事务处理和外键约束。'
    '数据库字符集采用utf8mb4编码，以全面支持中文及Emoji等Unicode字符。'
    '系统使用MySQL Connector/J 8.0.33驱动进行数据库连接。'
)

doc.add_heading('2.5 JWT身份认证', level=2)
add_paragraph(
    'JSON Web Token（JWT）是一种基于JSON的开放标准（RFC 7519），用于在各方之间安全传输信息。'
    'JWT由三部分组成：头部（Header）、载荷（Payload）和签名（Signature）。'
    '本系统采用jjwt 0.9.1库实现JWT身份认证机制。'
)
add_paragraph(
    '系统的认证流程为：用户登录成功后，服务端使用HS512算法生成JWT令牌，令牌中包含用户ID、用户名和角色信息，'
    '有效期设置为24小时（86400000毫秒）。客户端在后续请求中通过Authorization请求头携带Bearer Token，'
    '服务端通过AuthInterceptor拦截器验证令牌的有效性，并从令牌中提取用户信息注入到请求属性中。'
    '对于不需要认证的公开接口（如登录、注册、店铺列表等），拦截器配置了排除路径以放行请求。'
)

doc.add_heading('2.6 协同过滤推荐算法', level=2)
add_paragraph(
    '协同过滤（Collaborative Filtering）是推荐系统中最经典的算法之一，其核心思想是利用群体的行为来预测个体的偏好。'
    '本系统实现了基于用户的协同过滤推荐算法（User-Based CF），结合基于内容的推荐（Content-Based Recommendation）'
    '构建混合推荐策略。'
)
add_paragraph(
    '具体的推荐策略包含四个层次：'
    '（1）协同过滤推荐：通过分析用户行为记录（浏览、点击、收藏、下单、评价等），'
    '计算用户之间的相似度，找到与目标用户行为相似的用户群体，推荐这些相似用户喜欢但目标用户未接触的店铺和菜品；'
    '（2）基于内容的推荐：根据用户设置的口味偏好（辣度、甜度、酸度等）和偏好分类，'
    '推荐符合用户口味的店铺和菜品；'
    '（3）基于行为权重的推荐：对用户的各类行为赋予不同的权重分值'
    '（浏览1.0、点击2.0、收藏3.0、分享3.0、评价4.0、下单5.0），按加权得分排序推荐；'
    '（4）热门补充：当个性化推荐结果不足时，使用评分和销量排序的热门店铺/菜品进行补充。'
)

doc.add_heading('2.7 其他技术组件', level=2)

add_table_caption('表2-1 系统使用的其他技术组件')
make_three_line_table(
    ['组件名称', '版本', '用途说明'],
    [
        ['Lombok', '-', '通过注解自动生成getter/setter/toString等代码，减少样板代码'],
        ['Hutool', '5.8.18', 'Java工具类库，本系统主要使用其MD5加密工具（DigestUtil）'],
        ['FastJSON', '1.2.83', '阿里巴巴开源的JSON解析库，用于解析用户口味偏好JSON数据'],
        ['Knife4j', '3.0.3', '基于Swagger的API文档增强工具，提供接口文档自动生成功能'],
        ['ECharts', '5.4.3', 'Apache开源的数据可视化图表库，用于管理后台的数据可视化'],
        ['Axios', '1.6.0', '基于Promise的HTTP客户端，用于管理后台的网络请求'],
    ]
)

doc.add_page_break()

# ======================== 第三章 系统需求分析 ========================
doc.add_heading('第三章  系统需求分析', level=1)

doc.add_heading('3.1 功能性需求分析', level=2)

doc.add_heading('3.1.1 用户端功能需求', level=3)
add_paragraph(
    '用户端面向校园学生用户群体，主要功能需求包括：'
)
add_paragraph('（1）账户管理：支持用户名密码注册与登录，支持微信一键授权登录，支持个人信息编辑和口味偏好设置。')
add_paragraph('（2）美食发现：提供首页轮播Banner、美食分类导航、热门店铺展示、热门菜品展示、搜索功能，用户可通过多种方式发现美食。')
add_paragraph('（3）个性化推荐：系统基于用户行为数据和口味偏好，智能推荐符合用户喜好的店铺和菜品。')
add_paragraph('（4）店铺与菜品浏览：支持查看店铺详情（包含地址、营业时间、评分、菜品列表等），支持查看菜品详情（包含价格、描述、评分等）。')
add_paragraph('（5）订单管理：支持在线下单、查看订单列表和订单详情、跟踪订单状态。')
add_paragraph('（6）评价功能：支持对已完成订单进行多维度评价（口味、服务、环境），支持上传评价图片。')
add_paragraph('（7）收藏与分享：支持收藏/取消收藏店铺和菜品，支持将店铺和菜品分享给微信好友。')

doc.add_heading('3.1.2 商家端功能需求', level=3)
add_paragraph('商家端面向入驻平台的餐饮商家，主要功能需求包括：')
add_paragraph('（1）商家账户：支持商家注册（需审核）和登录，支持商家信息管理。')
add_paragraph('（2）店铺管理：支持创建和编辑店铺信息，包括名称、地址、营业时间、图片等。')
add_paragraph('（3）菜品管理：支持菜品的增加、编辑、上下架操作。')
add_paragraph('（4）订单处理：支持查看和处理用户订单，包括确认订单、完成订单等。')
add_paragraph('（5）评价管理：支持查看用户评价并进行回复。')
add_paragraph('（6）经营统计：提供订单数量、营业额等经营数据统计。')

doc.add_heading('3.1.3 管理员端功能需求', level=3)
add_paragraph('管理员端面向系统运营管理人员，主要功能需求包括：')
add_paragraph('（1）数据统计：展示系统整体运营数据，包括用户数、商家数、店铺数、菜品数、订单数、评价数等。')
add_paragraph('（2）数据可视化：通过图表展示分类分布、评分排行、销量排行、订单状态分布、用户画像等数据。')
add_paragraph('（3）用户管理：查看用户列表，支持按关键词搜索，可禁用/启用用户账号。')
add_paragraph('（4）商家管理：查看商家列表，审核新注册商家，管理商家状态。')
add_paragraph('（5）评价管理：审核管理用户评价内容，处理违规评价。')
add_paragraph('（6）违规管理：记录和管理用户/商家的违规行为，支持警告、禁用等处理措施。')

doc.add_heading('3.2 非功能性需求分析', level=2)

add_table_caption('表3-1 非功能性需求')
make_three_line_table(
    ['需求类别', '具体要求'],
    [
        ['性能需求', 'API接口响应时间不超过500ms；支持100+并发用户同时在线'],
        ['安全需求', '用户密码使用MD5加密存储；接口采用JWT认证；敏感操作需权限校验'],
        ['可用性需求', '系统可用性达到99%以上；支持断线重连和错误重试'],
        ['兼容性需求', '小程序支持iOS和Android双平台；管理后台支持主流浏览器（Chrome、Firefox、Edge）'],
        ['可维护性需求', '代码结构清晰，采用MVC分层架构；关键逻辑有注释说明'],
        ['数据需求', '数据库使用InnoDB引擎支持事务；时间字段自动填充；字符集支持中文和Emoji'],
    ]
)

doc.add_heading('3.3 系统用例分析', level=2)
add_paragraph(
    '根据上述功能需求分析，系统涉及三类参与者：普通用户（User）、商家（Merchant）和管理员（Admin）。'
    '各参与者的主要用例如下：'
)

add_table_caption('表3-2 系统用例列表')
make_three_line_table(
    ['参与者', '用例名称', '用例描述'],
    [
        ['用户', '注册登录', '通过用户名密码或微信授权完成身份认证'],
        ['用户', '浏览美食', '按分类、搜索关键词查看店铺和菜品信息'],
        ['用户', '获取推荐', '系统根据用户偏好智能推荐个性化美食'],
        ['用户', '下单消费', '选择菜品创建订单，跟踪订单状态'],
        ['用户', '发表评价', '对已消费订单进行评分和文字评价'],
        ['用户', '收藏分享', '收藏喜爱的店铺/菜品并分享给好友'],
        ['商家', '管理店铺', '创建和维护店铺信息及菜品列表'],
        ['商家', '处理订单', '接收、确认和完成用户订单'],
        ['商家', '回复评价', '对用户评价进行回复互动'],
        ['商家', '查看统计', '查看经营数据统计报表'],
        ['管理员', '系统监控', '查看系统运营数据和可视化统计'],
        ['管理员', '审核管理', '审核商家入驻申请和评价内容'],
        ['管理员', '违规处理', '对违规用户和商家进行警告或封禁'],
    ]
)

doc.add_page_break()

# ======================== 第四章 系统总体设计 ========================
doc.add_heading('第四章  系统总体设计', level=1)

doc.add_heading('4.1 系统架构设计', level=2)
add_paragraph(
    '本系统采用经典的前后端分离架构（B/S架构），整体分为三个层次：展示层、业务逻辑层和数据层。'
    '展示层由微信小程序客户端和Web管理后台组成，负责用户交互和页面展示；'
    '业务逻辑层由Spring Boot后端服务构成，以RESTful API的形式对外提供服务；'
    '数据层使用MySQL数据库进行数据持久化存储。'
)
add_paragraph(
    '系统架构的核心设计原则包括：'
    '（1）前后端分离：前端与后端通过HTTP/JSON协议通信，各自独立开发和部署；'
    '（2）RESTful API设计：后端接口遵循RESTful风格，使用统一的请求/响应格式（Result封装类）；'
    '（3）分层架构：后端代码按照Controller-Service-Mapper三层架构组织，职责分离清晰；'
    '（4）统一认证：采用JWT令牌机制实现统一的身份认证和权限控制。'
)

doc.add_heading('4.2 系统功能模块设计', level=2)
add_paragraph(
    '系统按照角色维度划分为用户端、商家端和管理员端三大模块，各模块进一步细分为若干子功能模块。'
)

add_table_caption('表4-1 系统功能模块划分')
make_three_line_table(
    ['一级模块', '二级模块', '核心功能'],
    [
        ['用户端', '首页模块', '轮播Banner、分类导航、搜索入口、热门推荐、个性化推荐'],
        ['用户端', '分类模块', '按美食分类浏览店铺和菜品'],
        ['用户端', '推荐模块', '个性化推荐店铺和菜品列表（需登录）'],
        ['用户端', '个人中心', '个人信息、口味偏好、收藏列表、订单列表、评价记录'],
        ['用户端', '店铺详情', '店铺信息、菜品列表、评价列表、收藏/分享'],
        ['用户端', '菜品详情', '菜品信息、评分、收藏/分享、加入订单'],
        ['商家端', '商家首页', '经营数据概览、快捷操作入口'],
        ['商家端', '店铺管理', '创建/编辑店铺、上传图片'],
        ['商家端', '菜品管理', '菜品CRUD、上下架管理'],
        ['商家端', '订单管理', '订单列表、订单处理（确认/完成/取消）'],
        ['商家端', '经营统计', '订单统计、营业额统计'],
        ['管理员端', '数据看板', '核心指标统计、多维度数据可视化图表'],
        ['管理员端', '用户管理', '用户列表查询、状态管理'],
        ['管理员端', '商家管理', '商家列表查询、审核、状态管理'],
        ['管理员端', '评价管理', '评价审核、违规标记'],
        ['管理员端', '违规管理', '违规记录查看、处理措施执行'],
    ]
)

doc.add_heading('4.3 系统技术架构', level=2)

add_table_caption('表4-2 系统技术架构各层组成')
make_three_line_table(
    ['架构层次', '技术选型', '说明'],
    [
        ['前端展示层', '微信小程序（WXML+WXSS+JS）', '用户端和商家端界面'],
        ['前端展示层', 'HTML5+CSS3+JavaScript+ECharts', '管理员后台Web界面'],
        ['网络通信层', 'HTTP/HTTPS + JSON', 'RESTful API通信协议'],
        ['认证安全层', 'JWT (jjwt) + AuthInterceptor', '身份认证和接口权限控制'],
        ['业务逻辑层', 'Spring Boot 2.7.18', 'Web框架和业务逻辑处理'],
        ['数据访问层', 'MyBatis-Plus 3.5.3.1', 'ORM映射和数据库操作'],
        ['数据存储层', 'MySQL 8.0', '关系型数据库存储'],
        ['文件存储', '本地文件系统', '图片等静态资源存储'],
    ]
)

doc.add_page_break()

# ======================== 第五章 数据库设计 ========================
doc.add_heading('第五章  数据库设计', level=1)

doc.add_heading('5.1 数据库概述', level=2)
add_paragraph(
    '本系统使用MySQL 8.0关系型数据库，数据库名为campus_food，字符集采用utf8mb4（支持Emoji等4字节Unicode字符），'
    '排序规则为utf8mb4_unicode_ci。所有数据表均采用InnoDB存储引擎，支持事务处理、行级锁和外键约束。'
    '数据库共设计14张数据表，涵盖用户管理、商家管理、店铺管理、菜品管理、订单管理、评价管理、收藏分享、'
    '行为记录和违规管理等核心业务领域。'
)

doc.add_heading('5.2 数据表结构设计', level=2)

# ----- 用户表 -----
doc.add_heading('5.2.1 用户表（user）', level=3)
add_paragraph('用户表用于存储系统普通用户的账户信息和个人资料，是系统最基础的核心数据表之一。')

add_table_caption('表5-1 用户表（user）结构')
make_three_line_table(
    ['字段名', '数据类型', '约束', '说明'],
    [
        ['id', 'BIGINT', 'PK, AUTO_INCREMENT', '用户ID，自增主键'],
        ['openid', 'VARCHAR(64)', 'UNIQUE', '微信openid，用于微信登录'],
        ['username', 'VARCHAR(50)', 'NOT NULL, UNIQUE', '登录用户名，唯一'],
        ['password', 'VARCHAR(100)', 'NOT NULL', '密码（MD5加密存储）'],
        ['nickname', 'VARCHAR(50)', '-', '用户昵称'],
        ['avatar', 'VARCHAR(255)', '-', '头像URL地址'],
        ['phone', 'VARCHAR(20)', '-', '手机号码'],
        ['email', 'VARCHAR(100)', '-', '邮箱地址'],
        ['gender', 'TINYINT', '默认0', '性别：0未知/1男/2女'],
        ['status', 'TINYINT', '默认1', '状态：0禁用/1正常'],
        ['taste_preference', 'VARCHAR(500)', '-', '口味偏好（JSON格式）'],
        ['create_time', 'DATETIME', '自动填充', '创建时间'],
        ['update_time', 'DATETIME', '自动更新', '更新时间'],
    ]
)

# ----- 商家表 -----
doc.add_heading('5.2.2 商家表（merchant）', level=3)
add_paragraph('商家表用于存储入驻平台的餐饮商家的账户信息和基本资料。')

add_table_caption('表5-2 商家表（merchant）结构')
make_three_line_table(
    ['字段名', '数据类型', '约束', '说明'],
    [
        ['id', 'BIGINT', 'PK, AUTO_INCREMENT', '商家ID，自增主键'],
        ['username', 'VARCHAR(50)', 'NOT NULL, UNIQUE', '登录用户名，唯一'],
        ['password', 'VARCHAR(100)', 'NOT NULL', '密码（MD5加密存储）'],
        ['name', 'VARCHAR(100)', 'NOT NULL', '商家名称'],
        ['contact_name', 'VARCHAR(50)', '-', '联系人姓名'],
        ['phone', 'VARCHAR(20)', '-', '联系电话'],
        ['email', 'VARCHAR(100)', '-', '邮箱地址'],
        ['avatar', 'VARCHAR(255)', '-', '商家头像URL'],
        ['status', 'TINYINT', '默认0', '状态：0待审核/1正常/2禁用'],
        ['violation_count', 'INT', '默认0', '累计违规次数'],
        ['create_time', 'DATETIME', '自动填充', '创建时间'],
        ['update_time', 'DATETIME', '自动更新', '更新时间'],
    ]
)

# ----- 管理员表 -----
doc.add_heading('5.2.3 管理员表（admin）', level=3)
add_paragraph('管理员表用于存储系统管理员的账户信息。')

add_table_caption('表5-3 管理员表（admin）结构')
make_three_line_table(
    ['字段名', '数据类型', '约束', '说明'],
    [
        ['id', 'BIGINT', 'PK, AUTO_INCREMENT', '管理员ID，自增主键'],
        ['username', 'VARCHAR(50)', 'NOT NULL, UNIQUE', '登录用户名，唯一'],
        ['password', 'VARCHAR(100)', 'NOT NULL', '密码（MD5加密存储）'],
        ['name', 'VARCHAR(50)', '-', '管理员姓名'],
        ['avatar', 'VARCHAR(255)', '-', '头像URL'],
        ['role', 'VARCHAR(20)', '默认admin', '角色标识'],
        ['create_time', 'DATETIME', '自动填充', '创建时间'],
        ['update_time', 'DATETIME', '自动更新', '更新时间'],
    ]
)

# ----- 分类表 -----
doc.add_heading('5.2.4 美食分类表（category）', level=3)
add_paragraph('美食分类表用于管理系统中的美食分类信息，支持分类名称、图标和排序。')

add_table_caption('表5-4 美食分类表（category）结构')
make_three_line_table(
    ['字段名', '数据类型', '约束', '说明'],
    [
        ['id', 'BIGINT', 'PK, AUTO_INCREMENT', '分类ID，自增主键'],
        ['name', 'VARCHAR(50)', 'NOT NULL', '分类名称（如中餐、西餐等）'],
        ['icon', 'VARCHAR(255)', '-', '分类图标路径'],
        ['sort_order', 'INT', '默认0', '排序序号，越小越靠前'],
        ['status', 'TINYINT', '默认1', '状态：0禁用/1正常'],
        ['create_time', 'DATETIME', '自动填充', '创建时间'],
    ]
)

# ----- 店铺表 -----
doc.add_heading('5.2.5 店铺表（shop）', level=3)
add_paragraph('店铺表存储商家创建的店铺信息，是系统的核心业务表之一，包含地理位置、评分、销量等运营数据。')

add_table_caption('表5-5 店铺表（shop）结构')
make_three_line_table(
    ['字段名', '数据类型', '约束', '说明'],
    [
        ['id', 'BIGINT', 'PK, AUTO_INCREMENT', '店铺ID，自增主键'],
        ['merchant_id', 'BIGINT', 'NOT NULL, INDEX', '所属商家ID'],
        ['name', 'VARCHAR(100)', 'NOT NULL', '店铺名称'],
        ['logo', 'VARCHAR(255)', '-', '店铺Logo图片URL'],
        ['images', 'TEXT', '-', '店铺图片（JSON数组格式）'],
        ['description', 'TEXT', '-', '店铺描述'],
        ['address', 'VARCHAR(255)', '-', '店铺详细地址'],
        ['longitude', 'DECIMAL(10,7)', '-', '经度坐标'],
        ['latitude', 'DECIMAL(10,7)', '-', '纬度坐标'],
        ['phone', 'VARCHAR(20)', '-', '联系电话'],
        ['business_hours', 'VARCHAR(100)', '-', '营业时间'],
        ['avg_price', 'DECIMAL(10,2)', '-', '人均消费金额（元）'],
        ['rating', 'DECIMAL(3,2)', '默认0.00', '综合评分（1.00-5.00）'],
        ['total_sales', 'INT', '默认0', '累计总销量'],
        ['category_id', 'BIGINT', 'INDEX', '所属美食分类ID'],
        ['tags', 'VARCHAR(255)', '-', '标签（逗号分隔）'],
        ['status', 'TINYINT', '默认1', '状态：0关闭/1营业'],
        ['create_time', 'DATETIME', '自动填充', '创建时间'],
        ['update_time', 'DATETIME', '自动更新', '更新时间'],
    ]
)

# ----- 菜品表 -----
doc.add_heading('5.2.6 菜品表（dish）', level=3)
add_paragraph('菜品表存储各店铺的菜品信息，包含价格、辣度、销量、评分等属性。')

add_table_caption('表5-6 菜品表（dish）结构')
make_three_line_table(
    ['字段名', '数据类型', '约束', '说明'],
    [
        ['id', 'BIGINT', 'PK, AUTO_INCREMENT', '菜品ID，自增主键'],
        ['shop_id', 'BIGINT', 'NOT NULL, INDEX', '所属店铺ID'],
        ['name', 'VARCHAR(100)', 'NOT NULL', '菜品名称'],
        ['image', 'VARCHAR(255)', '-', '菜品图片URL'],
        ['description', 'TEXT', '-', '菜品描述'],
        ['price', 'DECIMAL(10,2)', 'NOT NULL', '当前售价（元）'],
        ['original_price', 'DECIMAL(10,2)', '-', '原价（元）'],
        ['category_id', 'BIGINT', 'INDEX', '所属分类ID'],
        ['tags', 'VARCHAR(255)', '-', '标签（逗号分隔）'],
        ['spicy_level', 'TINYINT', '默认0', '辣度：0不辣/1微辣/2中辣/3特辣'],
        ['sales', 'INT', '默认0', '累计销量'],
        ['rating', 'DECIMAL(3,2)', '默认0.00', '综合评分'],
        ['status', 'TINYINT', '默认1', '状态：0下架/1上架'],
        ['sort_order', 'INT', '默认0', '排序序号'],
        ['create_time', 'DATETIME', '自动填充', '创建时间'],
        ['update_time', 'DATETIME', '自动更新', '更新时间'],
    ]
)

# ----- 订单表 -----
doc.add_heading('5.2.7 订单表（order_info）', level=3)
add_paragraph('订单表记录用户的消费订单信息，包含订单编号、金额、状态等核心业务数据。')

add_table_caption('表5-7 订单表（order_info）结构')
make_three_line_table(
    ['字段名', '数据类型', '约束', '说明'],
    [
        ['id', 'BIGINT', 'PK, AUTO_INCREMENT', '订单ID，自增主键'],
        ['order_no', 'VARCHAR(64)', 'NOT NULL, UNIQUE', '订单编号（UUID生成）'],
        ['user_id', 'BIGINT', 'NOT NULL, INDEX', '下单用户ID'],
        ['shop_id', 'BIGINT', 'NOT NULL, INDEX', '所属店铺ID'],
        ['merchant_id', 'BIGINT', 'NOT NULL, INDEX', '所属商家ID'],
        ['total_amount', 'DECIMAL(10,2)', 'NOT NULL', '订单总金额（元）'],
        ['status', 'TINYINT', '默认0', '状态：0待确认/1已确认/2已完成/3已取消'],
        ['remark', 'VARCHAR(255)', '-', '用户备注'],
        ['create_time', 'DATETIME', '自动填充', '下单时间'],
        ['update_time', 'DATETIME', '自动更新', '更新时间'],
    ]
)

# ----- 订单明细表 -----
doc.add_heading('5.2.8 订单明细表（order_item）', level=3)
add_paragraph('订单明细表记录每个订单中包含的菜品明细，与订单表构成一对多关系。')

add_table_caption('表5-8 订单明细表（order_item）结构')
make_three_line_table(
    ['字段名', '数据类型', '约束', '说明'],
    [
        ['id', 'BIGINT', 'PK, AUTO_INCREMENT', '明细ID，自增主键'],
        ['order_id', 'BIGINT', 'NOT NULL, INDEX', '所属订单ID'],
        ['dish_id', 'BIGINT', 'NOT NULL', '菜品ID'],
        ['dish_name', 'VARCHAR(100)', '-', '菜品名称（冗余存储）'],
        ['dish_image', 'VARCHAR(255)', '-', '菜品图片（冗余存储）'],
        ['price', 'DECIMAL(10,2)', 'NOT NULL', '下单时单价（元）'],
        ['quantity', 'INT', 'NOT NULL, 默认1', '购买数量'],
        ['amount', 'DECIMAL(10,2)', 'NOT NULL', '小计金额 = 单价 × 数量'],
    ]
)

# ----- 评价表 -----
doc.add_heading('5.2.9 评价表（review）', level=3)
add_paragraph('评价表记录用户对店铺和菜品的评价信息，支持多维度评分和图片上传。')

add_table_caption('表5-9 评价表（review）结构')
make_three_line_table(
    ['字段名', '数据类型', '约束', '说明'],
    [
        ['id', 'BIGINT', 'PK, AUTO_INCREMENT', '评价ID，自增主键'],
        ['user_id', 'BIGINT', 'NOT NULL, INDEX', '评价用户ID'],
        ['shop_id', 'BIGINT', 'NOT NULL, INDEX', '评价店铺ID'],
        ['dish_id', 'BIGINT', 'INDEX', '评价菜品ID（可选）'],
        ['order_id', 'BIGINT', '-', '关联订单ID（可选）'],
        ['content', 'TEXT', '-', '评价文字内容'],
        ['rating', 'TINYINT', 'NOT NULL', '综合评分（1-5）'],
        ['images', 'TEXT', '-', '评价图片（JSON数组格式）'],
        ['taste_rating', 'TINYINT', '-', '口味评分（1-5）'],
        ['service_rating', 'TINYINT', '-', '服务评分（1-5）'],
        ['environment_rating', 'TINYINT', '-', '环境评分（1-5）'],
        ['status', 'TINYINT', '默认1', '状态：0隐藏/1正常/2违规'],
        ['reply', 'TEXT', '-', '商家回复内容'],
        ['reply_time', 'DATETIME', '-', '商家回复时间'],
        ['create_time', 'DATETIME', '自动填充', '评价时间'],
    ]
)

# ----- 收藏店铺表 -----
doc.add_heading('5.2.10 收藏店铺表（favorite_shop）', level=3)
add_paragraph('收藏店铺表记录用户收藏的店铺信息，用户和店铺的组合具有唯一约束。')

add_table_caption('表5-10 收藏店铺表（favorite_shop）结构')
make_three_line_table(
    ['字段名', '数据类型', '约束', '说明'],
    [
        ['id', 'BIGINT', 'PK, AUTO_INCREMENT', 'ID，自增主键'],
        ['user_id', 'BIGINT', 'NOT NULL, INDEX', '用户ID'],
        ['shop_id', 'BIGINT', 'NOT NULL', '店铺ID'],
        ['create_time', 'DATETIME', '自动填充', '收藏时间'],
    ]
)

# ----- 收藏美食表 -----
doc.add_heading('5.2.11 收藏美食表（favorite_food）', level=3)
add_paragraph('收藏美食表记录用户收藏的菜品信息，用户和菜品的组合具有唯一约束。')

add_table_caption('表5-11 收藏美食表（favorite_food）结构')
make_three_line_table(
    ['字段名', '数据类型', '约束', '说明'],
    [
        ['id', 'BIGINT', 'PK, AUTO_INCREMENT', 'ID，自增主键'],
        ['user_id', 'BIGINT', 'NOT NULL, INDEX', '用户ID'],
        ['dish_id', 'BIGINT', 'NOT NULL', '菜品ID'],
        ['create_time', 'DATETIME', '自动填充', '收藏时间'],
    ]
)

# ----- 分享记录表 -----
doc.add_heading('5.2.12 分享记录表（share_record）', level=3)
add_paragraph('分享记录表记录用户的分享行为，用于统计分享数据和辅助推荐算法。')

add_table_caption('表5-12 分享记录表（share_record）结构')
make_three_line_table(
    ['字段名', '数据类型', '约束', '说明'],
    [
        ['id', 'BIGINT', 'PK, AUTO_INCREMENT', 'ID，自增主键'],
        ['user_id', 'BIGINT', 'NOT NULL, INDEX', '分享用户ID'],
        ['target_type', 'VARCHAR(20)', 'NOT NULL', '分享类型：shop/dish'],
        ['target_id', 'BIGINT', 'NOT NULL', '分享目标ID'],
        ['share_type', 'VARCHAR(20)', '-', '分享渠道'],
        ['create_time', 'DATETIME', '自动填充', '分享时间'],
    ]
)

# ----- 用户行为记录表 -----
doc.add_heading('5.2.13 用户行为记录表（user_behavior）', level=3)
add_paragraph('用户行为记录表是个性化推荐算法的核心数据源，记录用户的各类行为及其权重分值。')

add_table_caption('表5-13 用户行为记录表（user_behavior）结构')
make_three_line_table(
    ['字段名', '数据类型', '约束', '说明'],
    [
        ['id', 'BIGINT', 'PK, AUTO_INCREMENT', 'ID，自增主键'],
        ['user_id', 'BIGINT', 'NOT NULL, INDEX', '用户ID'],
        ['behavior_type', 'VARCHAR(20)', 'NOT NULL', '行为类型：view/click/order/review/favorite/share'],
        ['target_type', 'VARCHAR(20)', 'NOT NULL, INDEX', '目标类型：shop/dish/category'],
        ['target_id', 'BIGINT', 'NOT NULL, INDEX', '目标对象ID'],
        ['score', 'DECIMAL(5,2)', '默认1.00', '行为权重分值'],
        ['create_time', 'DATETIME', '自动填充', '记录时间'],
    ]
)

# ----- 违规记录表 -----
doc.add_heading('5.2.14 违规记录表（violation_record）', level=3)
add_paragraph('违规记录表用于记录和管理用户或商家的违规行为及处理措施。')

add_table_caption('表5-14 违规记录表（violation_record）结构')
make_three_line_table(
    ['字段名', '数据类型', '约束', '说明'],
    [
        ['id', 'BIGINT', 'PK, AUTO_INCREMENT', 'ID，自增主键'],
        ['target_type', 'VARCHAR(20)', 'NOT NULL, INDEX', '违规对象类型：user/merchant'],
        ['target_id', 'BIGINT', 'NOT NULL, INDEX', '违规对象ID'],
        ['reason', 'TEXT', 'NOT NULL', '违规原因描述'],
        ['action', 'VARCHAR(50)', 'NOT NULL', '处理措施：warn/disable/ban'],
        ['admin_id', 'BIGINT', '-', '处理管理员ID'],
        ['status', 'TINYINT', '默认1', '状态：0已撤销/1生效中'],
        ['create_time', 'DATETIME', '自动填充', '处理时间'],
        ['update_time', 'DATETIME', '自动更新', '更新时间'],
    ]
)

doc.add_heading('5.3 数据表关系说明', level=2)
add_paragraph(
    '本系统数据库中各表之间的关联关系如下：'
)
add_paragraph('（1）商家（merchant）与店铺（shop）：一对多关系，一个商家可创建多个店铺（通过merchant_id关联）。')
add_paragraph('（2）店铺（shop）与菜品（dish）：一对多关系，一个店铺包含多个菜品（通过shop_id关联）。')
add_paragraph('（3）店铺（shop）与分类（category）：多对一关系，多个店铺可属于同一分类（通过category_id关联）。')
add_paragraph('（4）菜品（dish）与分类（category）：多对一关系，多个菜品可属于同一分类（通过category_id关联）。')
add_paragraph('（5）用户（user）与订单（order_info）：一对多关系，一个用户可创建多个订单（通过user_id关联）。')
add_paragraph('（6）订单（order_info）与订单明细（order_item）：一对多关系，一个订单包含多个菜品明细（通过order_id关联）。')
add_paragraph('（7）用户（user）与评价（review）：一对多关系，一个用户可发表多条评价（通过user_id关联）。')
add_paragraph('（8）用户（user）与收藏（favorite_shop/favorite_food）：一对多关系，一个用户可收藏多个店铺和菜品。')
add_paragraph('（9）用户（user）与行为记录（user_behavior）：一对多关系，记录用户的所有行为数据用于推荐算法。')
add_paragraph('（10）管理员（admin）与违规记录（violation_record）：一对多关系，记录由哪位管理员执行的违规处理。')

doc.add_page_break()

# ======================== 第六章 系统详细设计与实现 ========================
doc.add_heading('第六章  系统详细设计与实现', level=1)

doc.add_heading('6.1 后端接口设计', level=2)
add_paragraph(
    '本系统后端采用RESTful API风格设计接口，所有接口以"/api"为统一前缀。'
    '接口返回统一使用Result封装类，包含code（状态码）、message（提示信息）和data（业务数据）三个字段。'
    '状态码200表示操作成功，401表示未认证，500表示服务器错误。'
)

add_table_caption('表6-1 系统主要API接口列表')
make_three_line_table(
    ['接口路径', '方法', '功能描述', '是否需认证'],
    [
        ['/api/user/login', 'POST', '用户登录', '否'],
        ['/api/user/register', 'POST', '用户注册', '否'],
        ['/api/user/wxLogin', 'POST', '微信登录', '否'],
        ['/api/user/info', 'GET', '获取用户信息', '是'],
        ['/api/user/update', 'PUT', '更新用户信息', '是'],
        ['/api/user/taste', 'PUT', '更新口味偏好', '是'],
        ['/api/user/page', 'GET', '用户分页列表（管理员）', '是'],
        ['/api/user/status/{id}', 'PUT', '更新用户状态（管理员）', '是'],
        ['/api/merchant/login', 'POST', '商家登录', '否'],
        ['/api/merchant/register', 'POST', '商家注册', '否'],
        ['/api/merchant/info', 'GET', '获取商家信息', '是'],
        ['/api/merchant/update', 'PUT', '更新商家信息', '是'],
        ['/api/merchant/page', 'GET', '商家分页列表', '是'],
        ['/api/merchant/status/{id}', 'PUT', '更新商家状态', '是'],
        ['/api/shop/list', 'GET', '店铺分页列表', '否'],
        ['/api/shop/detail/{id}', 'GET', '店铺详情', '否'],
        ['/api/shop/add', 'POST', '新增店铺', '是'],
        ['/api/shop/update', 'PUT', '更新店铺', '是'],
        ['/api/shop/merchant', 'GET', '获取商家店铺', '是'],
        ['/api/dish/list', 'GET', '菜品分页列表', '否'],
        ['/api/dish/detail/{id}', 'GET', '菜品详情', '否'],
        ['/api/dish/add', 'POST', '新增菜品', '是'],
        ['/api/dish/update', 'PUT', '更新菜品', '是'],
        ['/api/dish/shop/{shopId}', 'GET', '获取店铺菜品', '否'],
        ['/api/order/create', 'POST', '创建订单', '是'],
        ['/api/order/user/page', 'GET', '用户订单列表', '是'],
        ['/api/order/merchant/page', 'GET', '商家订单列表', '是'],
        ['/api/order/detail/{id}', 'GET', '订单详情', '是'],
        ['/api/order/status/{id}', 'PUT', '更新订单状态', '是'],
        ['/api/review/add', 'POST', '发表评价', '是'],
        ['/api/review/list', 'GET', '评价列表', '否'],
        ['/api/review/reply/{id}', 'PUT', '商家回复评价', '是'],
        ['/api/favorite/shop', 'POST', '收藏店铺', '是'],
        ['/api/favorite/food', 'POST', '收藏菜品', '是'],
        ['/api/recommend/shops', 'GET', '推荐店铺', '是'],
        ['/api/recommend/dishes', 'GET', '推荐菜品', '是'],
        ['/api/recommend/behavior', 'POST', '记录用户行为', '是'],
        ['/api/recommend/popular/shops', 'GET', '热门店铺', '否'],
        ['/api/recommend/popular/dishes', 'GET', '热门菜品', '否'],
        ['/api/admin/login', 'POST', '管理员登录', '否'],
        ['/api/admin/statistics', 'GET', '系统统计数据', '是'],
        ['/api/admin/visualization', 'GET', '可视化数据', '是'],
        ['/api/category/list', 'GET', '分类列表', '否'],
        ['/api/upload', 'POST', '文件上传', '否'],
        ['/api/violation/add', 'POST', '添加违规记录', '是'],
        ['/api/violation/page', 'GET', '违规记录列表', '是'],
    ]
)

doc.add_heading('6.2 用户模块实现', level=2)
add_paragraph(
    '用户模块是系统的基础模块，负责用户的注册、登录、信息管理和口味偏好设置。'
    '用户模块的后端实现由UserController和UserServiceImpl协同完成。'
)

doc.add_heading('6.2.1 用户注册', level=3)
add_paragraph(
    '用户注册功能接收用户名和密码等信息，首先检查用户名是否已存在，若不存在则将密码进行MD5加密后存入数据库。'
    '注册成功后返回用户信息（密码字段置空）。注册时用户状态默认设置为1（正常）。'
)
add_code_block('''public User register(User user) {
    User existing = this.getOne(new LambdaQueryWrapper<User>()
            .eq(User::getUsername, user.getUsername()));
    if (existing != null) throw new RuntimeException("用户名已存在");
    user.setPassword(DigestUtil.md5Hex(user.getPassword()));
    user.setStatus(1);
    this.save(user);
    user.setPassword(null);
    return user;
}''', '代码6-1 用户注册核心代码')

doc.add_heading('6.2.2 用户登录', level=3)
add_paragraph(
    '系统支持两种登录方式：用户名密码登录和微信授权登录。'
    '用户名密码登录：将用户输入的密码进行MD5加密后与数据库中的密码进行比对，验证通过后检查账户状态是否正常，'
    '最后使用JwtUtil生成JWT令牌，返回令牌和用户信息。'
    '微信授权登录：接收小程序端传来的code，调用微信接口获取openid，根据openid查询用户是否已存在，'
    '若不存在则自动创建新用户，最后生成JWT令牌返回。'
)
add_code_block('''public Map<String, Object> login(String username, String password) {
    String md5Password = DigestUtil.md5Hex(password);
    User user = this.getOne(new LambdaQueryWrapper<User>()
            .eq(User::getUsername, username)
            .eq(User::getPassword, md5Password));
    if (user == null) throw new RuntimeException("用户名或密码错误");
    if (user.getStatus() == 0) throw new RuntimeException("账号已被禁用");
    String token = jwtUtil.generateToken(user.getId(), user.getUsername(), "user");
    Map<String, Object> result = new HashMap<>();
    result.put("token", token); result.put("userInfo", user);
    return result;
}''', '代码6-2 用户登录核心代码')

doc.add_heading('6.2.3 口味偏好设置', level=3)
add_paragraph(
    '用户可以设置个人口味偏好，偏好数据以JSON格式存储在user表的taste_preference字段中。'
    'JSON结构包含辣度偏好（spicy）、甜度偏好（sweet）、酸度偏好（sour）和偏好分类（categories数组）。'
    '例如：{"spicy": 3, "sweet": 2, "sour": 1, "categories": [1, 5]}表示用户偏好辣食，'
    '喜欢中餐和火锅烧烤类别。该偏好数据将被推荐算法用于生成个性化推荐。'
)
add_code_block('''@PutMapping("/taste")
public Result<?> updateTastePreference(@RequestBody Map<String, String> params,
                                       HttpServletRequest request) {
    Long userId = (Long) request.getAttribute("userId");
    userService.updateTastePreference(userId, params.get("tastePreference"));
    return Result.success();
}''', '代码6-3 口味偏好设置接口')

doc.add_heading('6.3 商家模块实现', level=2)
add_paragraph(
    '商家模块提供商家的注册、登录、信息管理等功能。商家注册后默认状态为"待审核"（status=0），'
    '需要管理员在后台审核通过（status=1）后才能正常使用系统功能。'
    '商家登录成功后，JWT令牌中的角色标识为"merchant"。'
    '商家信息管理包括查看和编辑商家名称、联系人、联系电话等信息。'
    '管理员可以通过商家管理接口查看商家列表（支持分页和关键词搜索），并对商家进行审核、禁用或启用操作。'
)
add_code_block('''@PutMapping("/status/{id}")
public Result<?> updateMerchantStatus(@PathVariable Long id,
        @RequestBody Map<String, Integer> params) {
    merchantService.updateMerchantStatus(id, params.get("status"));
    return Result.success();
}''', '代码6-4 商家状态管理（审核/禁用）接口')

doc.add_heading('6.4 店铺与菜品模块实现', level=2)
doc.add_heading('6.4.1 店铺管理', level=3)
add_paragraph(
    '店铺模块支持商家创建和管理自己的店铺。创建店铺时需填写店铺名称、地址、营业时间、联系电话等基本信息，'
    '并可上传店铺Logo和环境图片。店铺列表查询支持按分类筛选（categoryId）、关键词搜索（keyword）和商家筛选（merchantId），'
    '使用MyBatis-Plus自定义分页查询实现。店铺详情接口会额外关联查询分类名称和商家名称信息。'
)
add_code_block('''public Shop getShopDetail(Long shopId) {
    Shop shop = this.getById(shopId);
    if (shop != null) {
        Category category = categoryMapper.selectById(shop.getCategoryId());
        if (category != null) shop.setCategoryName(category.getName());
        Merchant merchant = merchantMapper.selectById(shop.getMerchantId());
        if (merchant != null) shop.setMerchantName(merchant.getName());
    }
    return shop;
}''', '代码6-5 店铺详情查询核心代码')

doc.add_heading('6.4.2 菜品管理', level=3)
add_paragraph(
    '菜品模块支持商家在店铺下创建和管理菜品。每个菜品包含名称、价格、原价、描述、图片、辣度等级、分类、标签等信息。'
    '菜品列表支持按店铺（shopId）、分类（categoryId）和关键词（keyword）进行多维度筛选查询。'
    '商家可以对菜品进行上下架操作（更新status字段），下架的菜品将不会在用户端展示。'
)
add_code_block('''@GetMapping("/list")
public Result<?> getDishList(@RequestParam(defaultValue = "1") int page,
        @RequestParam(defaultValue = "10") int size,
        @RequestParam(required = false) Long shopId,
        @RequestParam(required = false) Long categoryId,
        @RequestParam(required = false) String keyword) {
    return Result.success(dishService.getDishPage(page, size,
            shopId, categoryId, keyword));
}''', '代码6-6 菜品列表多条件分页查询接口')

doc.add_heading('6.5 订单模块实现', level=2)
add_paragraph(
    '订单模块实现了完整的订单生命周期管理。创建订单时使用@Transactional注解保证事务一致性，'
    '订单编号通过UUID生成20位唯一字符串，初始状态为"待确认"（status=0）。'
    '订单创建过程：先保存订单主表（order_info），再遍历订单明细列表逐条保存（order_item），'
    '每条明细的金额（amount）由单价（price）乘以数量（quantity）计算得出。'
)
add_code_block('''@Transactional
public OrderInfo createOrder(OrderInfo order) {
    order.setOrderNo(UUID.randomUUID().toString().replace("-","").substring(0,20));
    order.setStatus(0);
    this.save(order);
    if (order.getItems() != null) {
        for (OrderItem item : order.getItems()) {
            item.setOrderId(order.getId());
            item.setAmount(item.getPrice().multiply(new BigDecimal(item.getQuantity())));
            orderItemMapper.insert(item);
        }
    }
    return order;
}''', '代码6-7 订单创建事务核心代码')
add_paragraph(
    '订单状态流转包含四个状态：待确认（0）→ 已确认（1）→ 已完成（2），或在任意阶段可转为已取消（3）。'
    '用户端和商家端分别可以查看各自的订单列表（按userId或merchantId分页查询），'
    '订单详情接口会关联查询订单明细列表。'
    '商家端还提供经营统计功能，统计总订单数、待处理订单数、已完成订单数和总营业额。'
)

doc.add_heading('6.6 评价模块实现', level=2)
add_paragraph(
    '评价模块支持用户对已消费的店铺和菜品进行多维度评价。评价内容包括文字评价（content）、'
    '综合评分（rating，1-5分）、口味评分（taste_rating）、服务评分（service_rating）、'
    '环境评分（environment_rating），以及可选的评价图片（images，以JSON数组格式存储多张图片URL）。'
)
add_code_block('''@PostMapping("/add")
public Result<?> addReview(@RequestBody Review review,
                           HttpServletRequest request) {
    Long userId = (Long) request.getAttribute("userId");
    review.setUserId(userId);
    return Result.success(reviewService.addReview(review));
}''', '代码6-8 发表评价接口')
add_paragraph(
    '商家可以对用户评价进行回复，回复内容存储在reply字段，并记录回复时间（reply_time）。'
    '管理员可以审核评价内容，将违规评价标记为违规状态（status=2）或隐藏（status=0）。'
    '评价列表支持按店铺、用户、状态等多条件分页查询。'
)

doc.add_heading('6.7 收藏与分享模块实现', level=2)
add_paragraph(
    '收藏功能支持用户收藏和取消收藏店铺与菜品。收藏店铺使用favorite_shop表存储，'
    '收藏菜品使用favorite_food表存储，两张表均设有联合唯一约束（user_id + shop_id/dish_id），'
    '确保同一用户不会重复收藏。系统提供收藏状态检查接口，用于前端展示收藏按钮的选中状态。'
)
add_code_block('''@PostMapping("/shop")
public Result<?> favoriteShop(@RequestBody Map<String, Long> params,
                              HttpServletRequest request) {
    Long userId = (Long) request.getAttribute("userId");
    favoriteService.favoriteShop(userId, params.get("shopId"));
    return Result.success();
}''', '代码6-9 收藏店铺接口')
add_paragraph(
    '分享功能记录用户的分享行为到share_record表中，包括分享目标类型（shop/dish）、'
    '目标ID和分享渠道等信息。分享行为同时会被推荐算法采集，作为用户偏好的重要参考依据。'
    '收藏和分享行为均通过调用推荐行为记录接口（/api/recommend/behavior）同步到用户行为记录表中。'
)

doc.add_heading('6.8 个性化推荐模块实现', level=2)
add_paragraph(
    '个性化推荐是本系统的核心亮点功能，由RecommendServiceImpl实现。推荐模块采用多策略混合推荐的架构，'
    '按优先级依次执行四个推荐策略，最终合并去重后返回推荐结果。'
)

doc.add_heading('6.8.1 行为权重设计', level=3)
add_paragraph('系统对用户的不同行为赋予不同的权重分值，权重越高表示用户对该目标的兴趣越强：')
add_code_block('''private static final Map<String, BigDecimal> BEHAVIOR_WEIGHTS =
    new HashMap<String, BigDecimal>() {{
        put("view", new BigDecimal("1.0"));
        put("click", new BigDecimal("2.0"));
        put("favorite", new BigDecimal("3.0"));
        put("share", new BigDecimal("3.0"));
        put("order", new BigDecimal("5.0"));
        put("review", new BigDecimal("4.0"));
    }};''', '代码6-10 行为权重定义')

add_table_caption('表6-2 用户行为权重表')
make_three_line_table(
    ['行为类型', '英文标识', '权重分值', '含义说明'],
    [
        ['浏览', 'view', '1.0', '用户浏览了店铺/菜品页面'],
        ['点击', 'click', '2.0', '用户主动点击了店铺/菜品'],
        ['收藏', 'favorite', '3.0', '用户收藏了店铺/菜品'],
        ['分享', 'share', '3.0', '用户分享了店铺/菜品'],
        ['评价', 'review', '4.0', '用户对店铺/菜品进行了评价'],
        ['下单', 'order', '5.0', '用户下单购买了菜品'],
    ]
)

doc.add_heading('6.8.2 协同过滤推荐策略', level=3)
add_paragraph(
    '协同过滤推荐是第一优先级的推荐策略。该策略分两步执行：'
    '第一步，查找与目标用户行为相似的用户群体（相似用户）。通过SQL聚合查询，'
    '统计与目标用户有共同行为目标（即对相同店铺/菜品有交互行为）的其他用户，'
    '按共同行为数量（similarity）降序排列，取前10名最相似用户。'
    '第二步，从相似用户群体中提取目标用户尚未接触的店铺/菜品，按行为加权总分排序，推荐给目标用户。'
)
add_code_block('''@Select("SELECT ub2.user_id, COUNT(*) as similarity FROM user_behavior ub1 "
  + "JOIN user_behavior ub2 ON ub1.target_type = ub2.target_type "
  + "AND ub1.target_id = ub2.target_id "
  + "WHERE ub1.user_id = #{userId} AND ub2.user_id != #{userId} "
  + "GROUP BY ub2.user_id ORDER BY similarity DESC LIMIT #{limit}")
List<Map<String, Object>> selectSimilarUsers(
    @Param("userId") Long userId, @Param("limit") int limit);''', '代码6-11 查找相似用户SQL（协同过滤核心）')

doc.add_heading('6.8.3 基于内容的推荐策略', level=3)
add_paragraph(
    '基于内容的推荐是第二优先级的推荐策略。该策略根据用户设置的口味偏好进行推荐：'
    '对于店铺推荐，从用户口味偏好JSON中提取偏好分类ID列表（categories），'
    '查询属于这些分类且处于营业状态的店铺，按评分降序排列。'
    '对于菜品推荐，提取用户的辣度偏好（spicy），查询辣度不超过用户偏好的菜品，'
    '同时根据偏好分类进行筛选，按评分降序排列。'
)
add_code_block('''private List<Long> getContentBasedShopRecommendations(Long userId, int limit) {
    User user = userMapper.selectById(userId);
    if (user == null || user.getTastePreference() == null) return new ArrayList<>();
    JSONObject pref = JSON.parseObject(user.getTastePreference());
    List<Integer> catIds = pref.getJSONArray("categories").toJavaList(Integer.class);
    if (catIds.isEmpty()) return new ArrayList<>();
    List<Shop> shops = shopMapper.selectList(new LambdaQueryWrapper<Shop>()
            .in(Shop::getCategoryId, catIds).eq(Shop::getStatus, 1)
            .orderByDesc(Shop::getRating).last("LIMIT " + limit));
    return shops.stream().map(Shop::getId).collect(Collectors.toList());
}''', '代码6-12 基于内容的店铺推荐核心代码')

doc.add_heading('6.8.4 行为权重排序策略', level=3)
add_paragraph(
    '第三优先级策略基于用户历史行为的加权得分进行推荐。通过SQL聚合查询用户行为记录表，'
    '按目标对象分组统计行为权重总分（SUM(score)），按总分降序排列，'
    '推荐用户最感兴趣的店铺/菜品。'
)
add_code_block('''@Select("SELECT target_id, SUM(score) as total_score "
  + "FROM user_behavior "
  + "WHERE user_id = #{userId} AND target_type = #{targetType} "
  + "GROUP BY target_id ORDER BY total_score DESC LIMIT #{limit}")
List<Map<String, Object>> selectUserPreferences(
    @Param("userId") Long userId,
    @Param("targetType") String targetType,
    @Param("limit") int limit);''', '代码6-13 用户行为加权排序SQL')

doc.add_heading('6.8.5 热门补充策略', level=3)
add_paragraph(
    '当以上三种个性化推荐策略产生的结果数量不足时，使用热门店铺/菜品进行补充。'
    '热门店铺按评分（rating）和总销量（total_sales）降序排列；'
    '热门菜品按销量（sales）和评分（rating）降序排列。'
    '对于未登录用户或新用户（无行为数据），系统直接展示热门推荐，确保推荐列表始终有内容展示。'
)
add_code_block('''public List<Shop> getPopularShops(int limit) {
    return shopMapper.selectList(new LambdaQueryWrapper<Shop>()
            .eq(Shop::getStatus, 1)
            .orderByDesc(Shop::getRating)
            .orderByDesc(Shop::getTotalSales)
            .last("LIMIT " + limit));
}''', '代码6-14 热门店铺查询核心代码')

doc.add_heading('6.9 管理员后台模块实现', level=2)
add_paragraph(
    '管理员后台由AdminController和AdminServiceImpl实现后端逻辑，前端使用纯HTML+CSS+JavaScript构建单页面应用，'
    '集成ECharts图表库进行数据可视化展示。'
)

doc.add_heading('6.9.1 数据统计', level=3)
add_paragraph(
    '系统统计接口（/api/admin/statistics）返回系统核心运营指标，包括：'
    '用户总数（userCount）、商家总数（merchantCount）、店铺总数（shopCount）、'
    '菜品总数（dishCount）、订单总数（orderCount）、评价总数（reviewCount）、'
    '待审核商家数（pendingMerchants）和违规评价数（violationReviews）。'
)
add_code_block('''public Map<String, Object> getStatistics() {
    Map<String, Object> stats = new HashMap<>();
    stats.put("userCount", userMapper.selectCount(null));
    stats.put("merchantCount", merchantMapper.selectCount(null));
    stats.put("shopCount", shopMapper.selectCount(null));
    stats.put("dishCount", dishMapper.selectCount(null));
    stats.put("orderCount", orderInfoMapper.selectCount(null));
    stats.put("reviewCount", reviewMapper.selectCount(null));
    stats.put("pendingMerchants", merchantMapper.selectCount(
        new LambdaQueryWrapper<Merchant>().eq(Merchant::getStatus, 0)));
    return stats;
}''', '代码6-15 管理后台数据统计核心代码')

doc.add_heading('6.9.2 数据可视化', level=3)
add_paragraph(
    '可视化接口（/api/admin/visualization）提供多维度的数据可视化数据，包括：'
    '（1）各分类店铺数量统计（饼图/柱状图）；'
    '（2）各分类菜品数量统计（饼图/柱状图）；'
    '（3）店铺评分Top10排行（条形图）；'
    '（4）菜品销量Top10排行（条形图）；'
    '（5）订单状态分布（环形图，待确认/已确认/已完成/已取消）；'
    '（6）用户性别分布（饼图）；'
    '（7）商家状态分布（饼图，待审核/正常/禁用）；'
    '（8）评价评分分布（柱状图，1-5星分布）。'
    '前端使用ECharts图表库将上述数据渲染为直观的可视化图表。'
)

doc.add_heading('6.10 系统安全与认证实现', level=2)

doc.add_heading('6.10.1 JWT认证机制', level=3)
add_paragraph(
    'JwtUtil工具类负责JWT令牌的生成和解析。生成令牌时，将用户ID（id）、用户名（username）和角色（role）'
    '写入Claims载荷，使用HS512签名算法和密钥进行签名，有效期为24小时。'
    '令牌验证时先解析Claims，再检查是否已过期。'
)
add_code_block('''public String generateToken(Long id, String username, String role) {
    Map<String, Object> claims = new HashMap<>();
    claims.put("id", id);
    claims.put("username", username);
    claims.put("role", role);
    return Jwts.builder().setClaims(claims).setSubject(username)
            .setIssuedAt(new Date())
            .setExpiration(new Date(System.currentTimeMillis() + expiration))
            .signWith(SignatureAlgorithm.HS512, secret).compact();
}''', '代码6-16 JWT令牌生成核心代码')

doc.add_heading('6.10.2 接口拦截器', level=3)
add_paragraph(
    'AuthInterceptor实现了HandlerInterceptor接口，在请求到达Controller之前进行拦截处理。'
    '拦截器首先放行OPTIONS预检请求（支持CORS跨域），然后从请求头中提取Authorization字段的Bearer Token，'
    '验证令牌有效性后将userId、username和role注入到request属性中，供后续业务逻辑使用。'
    '对于未携带令牌或令牌无效的请求，返回401状态码和错误提示。'
)
add_code_block('''public boolean preHandle(HttpServletRequest request,
        HttpServletResponse response, Object handler) throws IOException {
    if ("OPTIONS".equalsIgnoreCase(request.getMethod())) return true;
    String token = request.getHeader("Authorization");
    if (token != null && token.startsWith("Bearer "))
        token = token.substring(7);
    if (token == null || !jwtUtil.validateToken(token)) {
        response.setContentType("application/json;charset=UTF-8");
        response.getWriter().write(JSON.toJSONString(Result.error(401,"未登录")));
        return false;
    }
    request.setAttribute("userId", jwtUtil.getUserId(token));
    return true;
}''', '代码6-17 认证拦截器核心代码')

doc.add_heading('6.10.3 接口权限控制', level=3)
add_paragraph(
    'WebMvcConfig配置类通过addInterceptors方法注册认证拦截器，并配置需要认证和不需要认证的接口路径。'
    '以下接口被排除在认证范围之外（即不需要登录即可访问）：'
    '用户登录/注册接口、商家登录/注册接口、管理员登录接口、店铺列表和详情接口、'
    '菜品列表和详情接口、分类列表接口、评价列表接口、热门推荐接口、文件上传接口。'
    '其余所有"/api/**"路径下的接口均需要携带有效JWT令牌才能访问。'
)
add_code_block('''public void addInterceptors(InterceptorRegistry registry) {
    registry.addInterceptor(authInterceptor)
        .addPathPatterns("/api/**")
        .excludePathPatterns(
            "/api/user/login", "/api/user/register", "/api/user/wxLogin",
            "/api/merchant/login", "/api/merchant/register",
            "/api/admin/login", "/api/shop/list", "/api/shop/detail/**",
            "/api/dish/list", "/api/dish/detail/**", "/api/category/list",
            "/api/review/list", "/api/recommend/popular/**");
}''', '代码6-18 接口权限配置核心代码')

doc.add_heading('6.10.4 跨域配置', level=3)
add_paragraph(
    '系统通过WebMvcConfig配置CORS跨域策略，允许所有来源（allowedOriginPatterns = "*"）、'
    '所有HTTP方法（GET/POST/PUT/DELETE/OPTIONS）和所有请求头的跨域请求，'
    '并支持携带认证信息（allowCredentials = true），缓存预检结果3600秒。'
)

doc.add_heading('6.10.5 全局异常处理', level=3)
add_paragraph(
    'GlobalExceptionHandler使用@RestControllerAdvice注解实现全局异常拦截。'
    '对于RuntimeException（业务异常），直接返回异常消息作为错误提示；'
    '对于其他未预期的Exception，统一返回"系统异常，请稍后重试"的友好提示，'
    '避免将敏感的系统内部错误信息暴露给前端用户。'
)
add_code_block('''@RestControllerAdvice
public class GlobalExceptionHandler {
    @ExceptionHandler(RuntimeException.class)
    public Result<?> handleRuntimeException(RuntimeException e) {
        return Result.error(e.getMessage());
    }
    @ExceptionHandler(Exception.class)
    public Result<?> handleException(Exception e) {
        return Result.error("系统异常，请稍后重试");
    }
}''', '代码6-19 全局异常处理核心代码')

doc.add_page_break()

# ======================== 第七章 系统前端设计与实现 ========================
doc.add_heading('第七章  系统前端设计与实现', level=1)

doc.add_heading('7.1 小程序页面结构', level=2)
add_paragraph(
    '微信小程序前端共包含26个页面，采用TabBar + 子页面的导航结构。'
    'TabBar底部导航包含4个主页面：首页（index）、分类（category）、推荐（recommend）和我的（mine）。'
    '其余页面通过navigateTo进行跳转。'
)

add_table_caption('表7-1 小程序页面列表')
make_three_line_table(
    ['页面路径', '页面名称', '功能描述'],
    [
        ['pages/index/index', '首页', '轮播图、分类入口、热门推荐、个性化推荐'],
        ['pages/category/category', '分类页', '按美食分类浏览店铺和菜品'],
        ['pages/recommend/recommend', '推荐页', '个性化推荐店铺和菜品（需登录）'],
        ['pages/mine/mine', '个人中心', '用户信息、功能入口（订单/收藏/评价等）'],
        ['pages/login/login', '登录页', '用户名密码登录、微信登录'],
        ['pages/register/register', '注册页', '用户注册表单'],
        ['pages/shop/list/list', '店铺列表', '店铺分页列表，支持筛选和搜索'],
        ['pages/shop/detail/detail', '店铺详情', '店铺信息、菜品列表、评价'],
        ['pages/dish/detail/detail', '菜品详情', '菜品信息、评分、操作按钮'],
        ['pages/search/search', '搜索页', '关键词搜索店铺和菜品'],
        ['pages/taste/taste', '口味设置', '设置口味偏好（辣/甜/酸/分类）'],
        ['pages/order/create/create', '创建订单', '选择菜品、填写备注、确认下单'],
        ['pages/order/list/list', '订单列表', '用户订单列表，按状态筛选'],
        ['pages/review/add/add', '发表评价', '多维度评分、图片上传'],
        ['pages/review/list/list', '评价列表', '店铺/用户评价列表'],
        ['pages/favorite/favorite', '收藏列表', '收藏的店铺和菜品'],
        ['pages/merchant/login/login', '商家登录', '商家端登录入口'],
        ['pages/merchant/register/register', '商家注册', '商家注册表单'],
        ['pages/merchant/index/index', '商家首页', '经营数据、快捷操作'],
        ['pages/merchant/shop/edit/edit', '店铺编辑', '创建/编辑店铺信息'],
        ['pages/merchant/dish/list/list', '菜品列表', '商家菜品管理列表'],
        ['pages/merchant/dish/edit/edit', '菜品编辑', '创建/编辑菜品信息'],
        ['pages/merchant/order/list/list', '商家订单', '商家订单管理列表'],
        ['pages/merchant/stats/stats', '经营统计', '订单和营业额统计'],
        ['pages/merchant/info/info', '商家信息', '商家个人信息管理'],
        ['pages/merchant/review/list/list', '评价管理', '查看和回复用户评价'],
    ]
)

doc.add_heading('7.2 网络请求封装', level=2)
add_paragraph(
    '小程序前端通过utils/request.js封装了统一的网络请求工具。'
    '该模块基于wx.request API封装了Promise风格的请求方法，提供get、post、put、del四种快捷方法，'
    '以及uploadFile文件上传方法。请求工具自动处理以下通用逻辑：'
)
add_code_block('''const request = (options) => {
  return new Promise((resolve, reject) => {
    const { url, method = 'GET', data = {}, loading = true } = options
    const header = { 'Content-Type': 'application/json' }
    if (app.globalData.token)
      header['Authorization'] = 'Bearer ' + app.globalData.token
    wx.request({
      url: app.globalData.baseUrl + url,
      method, data, header,
      success(res) { res.data.code===200 ? resolve(res.data) : reject(res.data) },
      fail(err) { wx.showToast({title:'网络请求失败',icon:'none'}); reject(err) }
    })
  })
}''', '代码7-1 小程序网络请求封装核心代码')
add_paragraph('（1）自动从全局数据（globalData）中读取JWT令牌，添加到请求头的Authorization字段。')
add_paragraph('（2）统一的加载动画管理（loading参数控制是否显示加载提示）。')
add_paragraph('（3）统一的响应码处理：200成功、401未认证（自动清除登录信息并跳转登录页）、其他错误码显示提示。')
add_paragraph('（4）统一的网络异常处理，显示"网络请求失败"提示。')

doc.add_heading('7.3 管理后台前端', level=2)
add_paragraph(
    '管理员后台采用单页面应用（SPA）设计，以纯HTML+CSS+JavaScript实现，部署在Spring Boot的静态资源目录下。'
    '后台页面使用侧边栏导航（sidebar），包含数据看板、用户管理、商家管理、店铺管理、评价管理和违规管理等菜单。'
    '数据看板页面集成ECharts图表库，展示系统核心统计指标和多维度可视化图表。'
    '各管理页面均使用Axios进行API调用，支持数据表格展示、搜索筛选和状态管理操作。'
    '页面整体采用橙色（#ff6b35）为主题色，界面风格简洁现代。'
)
add_code_block('''App({
  globalData: {
    baseUrl: 'http://localhost:8080',
    userInfo: null, token: null, role: null
  },
  onLaunch() {
    const token = wx.getStorageSync('token');
    const userInfo = wx.getStorageSync('userInfo');
    if (token) {
      this.globalData.token = token;
      this.globalData.userInfo = userInfo;
    }
  }
})''', '代码7-2 小程序全局入口（App.js）核心代码')

doc.add_page_break()

# ======================== 第八章 系统部署与运行 ========================
doc.add_heading('第八章  系统部署与运行', level=1)

doc.add_heading('8.1 开发环境', level=2)

add_table_caption('表8-1 开发环境配置')
make_three_line_table(
    ['环境项目', '版本/工具', '说明'],
    [
        ['操作系统', 'Windows 10/11', '开发操作系统'],
        ['JDK', 'Java 17', '后端运行环境'],
        ['Maven', 'Apache Maven 3.9.6', '后端项目构建工具'],
        ['MySQL', 'MySQL 8.0', '数据库服务'],
        ['IDE', 'IntelliJ IDEA', '后端开发工具'],
        ['微信开发者工具', '最新稳定版', '小程序开发调试工具'],
        ['Node.js', '16+', '前端构建环境（可选）'],
        ['浏览器', 'Chrome', '管理后台调试'],
    ]
)

doc.add_heading('8.2 部署步骤', level=2)

doc.add_heading('8.2.1 数据库部署', level=3)
add_paragraph('（1）安装并启动MySQL 8.0服务。')
add_paragraph('（2）使用数据库客户端工具（如Navicat或MySQL Workbench）连接MySQL服务。')
add_paragraph('（3）执行sql/init.sql脚本，自动创建campus_food数据库及全部14张数据表，并插入初始测试数据。')
add_paragraph('（4）确认数据库连接信息与application.yml配置一致（默认用户名root，密码root，端口3306）。')

doc.add_heading('8.2.2 后端部署', level=3)
add_paragraph('（1）确保已安装JDK 17和Maven 3.9.6。')
add_paragraph('（2）进入backend目录，修改application.yml中的数据库连接信息（如需要）。')
add_paragraph('（3）执行mvn clean package命令打包项目，生成target目录下的jar文件。')
add_paragraph('（4）执行java -jar target/food-recommendation-1.0.0.jar启动后端服务。')
add_paragraph('（5）服务启动后，后端API默认运行在http://localhost:8080。')
add_paragraph('（6）管理后台访问地址为http://localhost:8080/admin/index.html，默认管理员账号admin，密码admin123。')

doc.add_heading('8.2.3 小程序部署', level=3)
add_paragraph('（1）下载并安装微信开发者工具。')
add_paragraph('（2）使用微信开发者工具打开miniprogram目录。')
add_paragraph('（3）确认app.js中的baseUrl配置指向正确的后端服务地址（默认http://localhost:8080）。')
add_paragraph('（4）在微信开发者工具的"详情" → "本地设置"中勾选"不校验合法域名"（开发阶段）。')
add_paragraph('（5）编译运行小程序，即可在模拟器或真机中预览和测试。')
add_paragraph('（6）测试用户账号：user01/123456；测试商家账号：merchant01/123456。')

doc.add_page_break()

# ======================== 第九章 总结与展望 ========================
doc.add_heading('第九章  总结与展望', level=1)

doc.add_heading('9.1 项目总结', level=2)
add_paragraph(
    '本项目设计并实现了一套基于微信小程序的校园周边美食推荐系统，系统采用Spring Boot + MyBatis-Plus + MySQL'
    '构建后端服务，微信小程序构建用户端和商家端，HTML+ECharts构建管理后台，实现了完整的三端协同业务生态。'
)
add_paragraph(
    '系统主要技术成果包括：'
    '（1）实现了基于协同过滤和内容推荐的混合推荐算法，通过六种用户行为（浏览、点击、收藏、分享、评价、下单）'
    '的加权分析，为用户提供个性化的美食推荐服务。'
    '（2）构建了涵盖用户管理、商家管理、店铺管理、菜品管理、订单管理、评价管理、收藏分享、违规管理等'
    '14张数据表的完整数据库模型。'
    '（3）设计了46个RESTful API接口，涵盖系统全部业务功能，接口设计规范统一。'
    '（4）实现了基于JWT的统一身份认证机制和接口权限控制，保障系统安全性。'
    '（5）管理后台集成了ECharts数据可视化，提供8种维度的数据图表展示。'
)

doc.add_heading('9.2 不足与展望', level=2)
add_paragraph(
    '尽管系统已实现了核心功能，但仍存在一些可以改进的方面：'
)
add_paragraph(
    '（1）推荐算法优化：当前的协同过滤算法在数据稀疏时效果有限，'
    '未来可以引入矩阵分解（SVD）、深度学习推荐模型（如DeepFM、Wide & Deep）等更先进的推荐算法，'
    '提升推荐的准确性和多样性。'
)
add_paragraph(
    '（2）实时通信：当前系统缺乏实时消息推送能力，未来可引入WebSocket技术，'
    '实现订单状态实时推送、商家即时接单等功能。'
)
add_paragraph(
    '（3）支付集成：当前订单功能为模拟下单，未接入真实支付系统，'
    '未来可集成微信支付SDK，实现完整的在线支付流程。'
)
add_paragraph(
    '（4）地图服务：虽然数据库已预留了经纬度字段，但前端尚未集成腾讯地图SDK进行位置展示和导航功能，'
    '未来可加入基于位置的店铺推荐和路线导航功能。'
)
add_paragraph(
    '（5）性能优化：对于热门数据（如分类列表、热门推荐）可引入Redis缓存层，'
    '减少数据库查询压力，提升系统响应速度。'
)
add_paragraph(
    '（6）安全增强：密码加密可从MD5升级为BCrypt等更安全的算法；'
    '增加接口限流、参数校验等安全防护措施。'
)

# ======================== 保存文档 ========================
output_dir = r'f:\26毕设单\美食推荐小程序\技术文档'
os.makedirs(output_dir, exist_ok=True)
output_path = os.path.join(output_dir, '基于微信小程序的校园周边美食推荐系统_技术文档.docx')
doc.save(output_path)
print(f'文档生成完成: {output_path}')
